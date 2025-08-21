#!/usr/bin/env python3
"""
创建测试文档：包含图像的PDF和DOCX文件
"""

import os
from pathlib import Path

def create_test_pdf():
    """创建包含图像的测试PDF"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.utils import ImageReader
        from PIL import Image, ImageDraw
        import io
        
        # 创建一个简单的测试图像
        def create_test_image(text, size=(200, 100), color='lightblue'):
            img = Image.new('RGB', size, color=color)
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), text, fill='black')
            return img
        
        # 创建PDF
        pdf_path = "test_document.pdf"
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        
        # 第一页
        c.drawString(100, height - 100, "DocuGenius 测试文档")
        c.drawString(100, height - 130, "这是第一页的内容。")
        c.drawString(100, height - 160, "下面应该有一张图片：")
        
        # 添加第一张图片
        img1 = create_test_image("测试图片 1", color='lightblue')
        img1_bytes = io.BytesIO()
        img1.save(img1_bytes, format='PNG')
        img1_bytes.seek(0)
        c.drawImage(ImageReader(img1_bytes), 100, height - 300, width=200, height=100)
        
        c.drawString(100, height - 330, "这是图片后面的文字。")
        c.showPage()
        
        # 第二页
        c.drawString(100, height - 100, "第二页内容")
        c.drawString(100, height - 130, "这里有另一张图片：")
        
        # 添加第二张图片
        img2 = create_test_image("测试图片 2", color='lightgreen')
        img2_bytes = io.BytesIO()
        img2.save(img2_bytes, format='PNG')
        img2_bytes.seek(0)
        c.drawImage(ImageReader(img2_bytes), 100, height - 300, width=200, height=100)
        
        c.drawString(100, height - 330, "第二页结束。")
        c.save()
        
        print(f"✅ 创建测试PDF: {pdf_path}")
        return pdf_path
        
    except ImportError as e:
        print(f"❌ 创建PDF需要安装依赖: pip install reportlab pillow")
        print(f"错误: {e}")
        return None
    except Exception as e:
        print(f"❌ 创建PDF失败: {e}")
        return None

def create_test_docx():
    """创建包含图像的测试DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image, ImageDraw
        import io
        
        # 创建测试图像
        def create_test_image(text, size=(300, 150), color='lightcoral'):
            img = Image.new('RGB', size, color=color)
            draw = ImageDraw.Draw(img)
            draw.text((10, 60), text, fill='black')
            return img
        
        # 创建DOCX文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('DocuGenius 测试文档', 0)
        
        # 添加段落
        doc.add_paragraph('这是一个测试文档，用于验证DocuGenius的图像提取功能。')
        doc.add_paragraph('下面是第一张图片：')
        
        # 添加第一张图片
        img1 = create_test_image("DOCX 测试图片 1", color='lightcoral')
        img1_bytes = io.BytesIO()
        img1.save(img1_bytes, format='PNG')
        img1_bytes.seek(0)
        doc.add_picture(img1_bytes, width=Inches(3))
        
        # 添加更多内容
        doc.add_paragraph('这是第一张图片后面的内容。')
        doc.add_paragraph('现在添加第二张图片：')
        
        # 添加第二张图片
        img2 = create_test_image("DOCX 测试图片 2", color='lightyellow')
        img2_bytes = io.BytesIO()
        img2.save(img2_bytes, format='PNG')
        img2_bytes.seek(0)
        doc.add_picture(img2_bytes, width=Inches(2.5))
        
        # 添加结尾
        doc.add_paragraph('这是文档的结尾。')
        
        # 保存文档
        docx_path = "test_document.docx"
        doc.save(docx_path)
        
        print(f"✅ 创建测试DOCX: {docx_path}")
        return docx_path
        
    except ImportError as e:
        print(f"❌ 创建DOCX需要安装依赖: pip install python-docx pillow")
        print(f"错误: {e}")
        return None
    except Exception as e:
        print(f"❌ 创建DOCX失败: {e}")
        return None

def main():
    print("🏗️  创建测试文档")
    print("=" * 30)
    
    # 创建测试文档
    pdf_path = create_test_pdf()
    docx_path = create_test_docx()
    
    created_files = []
    if pdf_path:
        created_files.append(pdf_path)
    if docx_path:
        created_files.append(docx_path)
    
    if created_files:
        print(f"\n✅ 成功创建 {len(created_files)} 个测试文档:")
        for file_path in created_files:
            print(f"  📄 {file_path}")
        print(f"\n💡 现在可以运行测试脚本:")
        print(f"   python test_image_extraction.py")
    else:
        print(f"\n❌ 未能创建测试文档")
        print(f"请安装必要的依赖:")
        print(f"  pip install reportlab python-docx pillow")

if __name__ == "__main__":
    main()
